import os
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass
import chardet
from pypdf import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown


@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]
    source: str


class DocumentLoader:
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._load_pdf,
            '.txt': self._load_text,
            '.md': self._load_text,
            '.html': self._load_html,
            '.docx': self._load_docx,
        }
    
    def load_documents(self, directory_path: str) -> List[Document]:
        documents = []
        path = Path(directory_path)
        
        if not path.exists():
            raise ValueError(f"Directory {directory_path} does not exist")
        
        for file_path in path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc = self.load_document(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")
        
        return documents
    
    def load_document(self, file_path: str) -> Document:
        path = Path(file_path)
        
        if not path.exists():
            raise ValueError(f"File {file_path} does not exist")
        
        extension = path.suffix.lower()
        
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}")
        
        loader_func = self.supported_extensions[extension]
        content = loader_func(file_path)
        
        metadata = {
            'source': str(path),
            'file_name': path.name,
            'file_type': extension,
            'file_size': path.stat().st_size,
        }
        
        return Document(content=content, metadata=metadata, source=str(path))
    
    def _load_pdf(self, file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        
        return '\n'.join(text)
    
    def _load_text(self, file_path: str) -> str:
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
        
        with open(file_path, 'r', encoding=encoding) as file:
            return file.read()
    
    def _load_html(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            
            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split('  '))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
    
    def _load_docx(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text:
                full_text.append(paragraph.text)
        
        return '\n'.join(full_text)